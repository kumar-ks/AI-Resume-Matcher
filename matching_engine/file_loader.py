"""
File Loader Module
==================

Extracts text content from PDF, DOCX, and TXT files.
Supports loading individual files or scanning entire directories.

CALL HIERARCHY:
    run.py → load_jd() / load_resumes()
        → extract_text(file_path)          # Public API: routes to format-specific reader
            → _read_txt(path)              # Plain text files
            → _read_pdf(path)              # PDF files (PyPDF2 → pdfplumber → OCR fallback)
                → _ocr_pdf(path)           # OCR for scanned/image-based PDFs
            → _read_docx(path)             # DOCX files (paragraphs + tables + text boxes + hyperlinks)
        → load_files_from_directory(dir)   # Batch load all supported files from a folder

SUPPORTED FORMATS:
    .pdf  - Text-based PDFs (PyPDF2/pdfplumber) and scanned PDFs (OCR via tesseract)
    .docx - Microsoft Word documents (paragraphs, tables, text boxes, hyperlinks)
    .doc  - Legacy Word documents (same handler as .docx, limited support)
    .txt  - Plain text files (UTF-8)

DEPENDENCIES:
    Required: PyPDF2, python-docx
    Optional: pdfplumber (complex PDF layouts), pytesseract + pdf2image (scanned PDFs)
"""

import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────────────────────


def extract_text(file_path: str | Path) -> str:
    """
    Extract text content from a file (PDF, DOCX, or TXT).

    Called by: run.py → load_jd(), load_resumes()
    Calls: _read_txt(), _read_pdf(), or _read_docx() based on file extension

    Args:
        file_path: Path to the file to extract text from

    Returns:
        Extracted text content as a string (may be empty for scanned PDFs without OCR)

    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file does not exist
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    ext = path.suffix.lower()
    logger.debug(f"Extracting text from: {path.name} (type: {ext})")

    if ext == ".txt":
        return _read_txt(path)
    elif ext == ".pdf":
        return _read_pdf(path)
    elif ext in (".docx", ".doc"):
        return _read_docx(path)
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )


def load_files_from_directory(directory: str | Path) -> list[dict]:
    """
    Load all supported files from a directory.

    Called by: run.py → load_jd(), load_resumes()
    Calls: extract_text() for each supported file found

    Args:
        directory: Path to directory containing files

    Returns:
        List of dicts with keys:
            - 'filename': Name of the file (e.g., "resume.pdf")
            - 'path': Full path to the file
            - 'text': Extracted text content

    Raises:
        FileNotFoundError: If directory does not exist
        ValueError: If path is not a directory
    """
    dir_path = Path(directory)

    if not dir_path.exists():
        raise FileNotFoundError(f"Directory not found: {dir_path}")

    if not dir_path.is_dir():
        raise ValueError(f"Not a directory: {dir_path}")

    results = []
    # Sort files alphabetically for consistent ordering across runs
    for file_path in sorted(dir_path.iterdir()):
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            try:
                text = extract_text(file_path)
                results.append({
                    "filename": file_path.name,
                    "path": str(file_path),
                    "text": text,
                })
                logger.info(f"Loaded: {file_path.name} ({len(text)} chars)")
            except Exception as e:
                logger.warning(f"Failed to load {file_path.name}: {e}")

    logger.info(f"Loaded {len(results)} files from {dir_path}")
    return results


# ─────────────────────────────────────────────────────────────────────────────
# PRIVATE: FORMAT-SPECIFIC READERS
# ─────────────────────────────────────────────────────────────────────────────


def _read_txt(path: Path) -> str:
    """
    Read plain text file with UTF-8 encoding.

    Called by: extract_text() when file extension is .txt
    """
    logger.debug(f"Reading TXT file: {path.name}")
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    """
    Extract text from PDF using a multi-strategy approach.

    Called by: extract_text() when file extension is .pdf
    Calls: _ocr_pdf() as final fallback for scanned documents

    Strategy (in order):
        1. PyPDF2 - Fast, works for most text-based PDFs
        2. pdfplumber - Better for complex layouts, tables, multi-column
        3. OCR (pytesseract) - For scanned/image-based PDFs

    Returns:
        Extracted text, or empty string if all strategies fail.
    """
    text = ""

    # ── Strategy 1: PyPDF2 (fastest, handles standard text PDFs) ──────────────
    try:
        from PyPDF2 import PdfReader

        logger.debug(f"Trying PyPDF2 for: {path.name}")
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)
        text = "\n".join(pages)
        if text.strip():
            logger.debug(f"PyPDF2 succeeded: {len(text)} chars from {path.name}")
    except ImportError:
        logger.debug("PyPDF2 not installed, skipping")
    except Exception as e:
        logger.debug(f"PyPDF2 failed for {path.name}: {e}")

    # ── Strategy 2: pdfplumber (better for complex layouts) ───────────────────
    if not text.strip():
        try:
            import pdfplumber

            logger.debug(f"Trying pdfplumber for: {path.name}")
            with pdfplumber.open(str(path)) as pdf:
                pages = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
                text = "\n".join(pages)
            if text.strip():
                logger.debug(f"pdfplumber succeeded: {len(text)} chars from {path.name}")
        except ImportError:
            logger.debug("pdfplumber not installed, skipping")
        except Exception as e:
            logger.debug(f"pdfplumber failed for {path.name}: {e}")

    # ── Strategy 3: OCR (for scanned/image-based PDFs) ────────────────────────
    if not text.strip():
        logger.debug(f"Text extraction failed, attempting OCR for: {path.name}")
        text = _ocr_pdf(path)

    # ── Final check: warn if nothing was extracted ────────────────────────────
    if not text.strip():
        logger.warning(
            f"Could not extract text from {path.name}. "
            "The PDF may be image-based (scanned). "
            "Install tesseract for OCR: brew install tesseract && pip install pytesseract pdf2image. "
            "Or convert the PDF to DOCX/TXT manually."
        )

    return text


def _ocr_pdf(path: Path) -> str:
    """
    Attempt OCR (Optical Character Recognition) on a PDF.

    Called by: _read_pdf() when PyPDF2 and pdfplumber both fail

    Requires system packages:
        - tesseract: brew install tesseract
        - poppler: brew install poppler (provides pdftoppm for pdf2image)

    Requires Python packages:
        - pytesseract: pip install pytesseract
        - pdf2image: pip install pdf2image

    Process:
        1. Convert each PDF page to an image (via pdf2image/poppler)
        2. Run tesseract OCR on each image
        3. Concatenate all page texts

    Returns:
        OCR-extracted text, or empty string if OCR is unavailable/fails.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_path

        logger.debug(f"Running OCR on: {path.name}")
        images = convert_from_path(str(path))
        pages = []
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img)
            if page_text.strip():
                pages.append(page_text)
                logger.debug(f"  OCR page {i+1}: {len(page_text)} chars")
        result = "\n".join(pages)
        logger.debug(f"OCR complete: {len(result)} total chars from {path.name}")
        return result
    except ImportError:
        logger.debug(
            "OCR dependencies not available (pytesseract, pdf2image). "
            "Install with: brew install tesseract poppler && pip install pytesseract pdf2image"
        )
        return ""
    except Exception as e:
        logger.debug(f"OCR failed for {path.name}: {e}")
        return ""


def _read_docx(path: Path) -> str:
    """
    Extract text from DOCX files, including all content containers.

    Called by: extract_text() when file extension is .docx or .doc

    Modern DOCX resumes store content in various containers:
        - Hyperlinks: Email/phone stored as mailto:/tel: links (not visible text)
        - Text boxes (w:txbxContent): Contact info, headers, sidebars
        - Headers/Footers: LinkedIn URLs, page numbers
        - Paragraphs: Main body text
        - Tables: Structured layouts (very common in resumes)

    Extraction order (matches typical resume layout priority):
        1. Hyperlinks (mailto:, tel:) → Contact info
        2. Text boxes → Contact bar, sidebar content
        3. Document headers → Header content
        4. Paragraphs → Main body text
        5. Tables → Structured content (skills grids, experience tables)
        6. Document footers → Footer content

    Returns:
        All extracted text joined by newlines.
    """
    try:
        import docx
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX support. "
            "Install with: pip install python-docx"
        )

    doc = docx.Document(str(path))
    parts = []

    # ── 1. Extract hyperlinks (email/phone often stored ONLY as hyperlinks) ───
    # In many DOCX templates, contact info is clickable but not in paragraph text
    hyperlink_texts = []
    for rel in doc.part.rels.values():
        target = str(rel.target_ref)
        if target.startswith("mailto:"):
            hyperlink_texts.append(target.replace("mailto:", ""))
        elif target.startswith("tel:"):
            hyperlink_texts.append(target.replace("tel:", ""))
    if hyperlink_texts:
        # Deduplicate (same link may appear multiple times)
        unique_links = list(dict.fromkeys(hyperlink_texts))
        parts.append(" | ".join(unique_links))
        logger.debug(f"  Extracted {len(unique_links)} hyperlinks from {path.name}")

    # ── 2. Extract from text boxes (w:txbxContent in OOXML) ───────────────────
    # Modern resume templates use text boxes for contact bars, sidebars, etc.
    # These are NOT accessible via doc.paragraphs or doc.tables.
    body = doc.element.body
    seen_textbox_content = set()
    for tb in body.findall('.//' + qn('w:txbxContent')):
        # Collect all w:t (text run) elements within this text box
        texts = []
        for t_elem in tb.iter(qn('w:t')):
            if t_elem.text:
                texts.append(t_elem.text)
        tb_text = " ".join(texts).strip()
        # Deduplicate: text boxes often appear twice in XML (drawing + fallback)
        if tb_text and tb_text not in seen_textbox_content:
            seen_textbox_content.add(tb_text)
            parts.append(tb_text)
    if seen_textbox_content:
        logger.debug(f"  Extracted {len(seen_textbox_content)} text boxes from {path.name}")

    # ── 3. Extract from document headers ──────────────────────────────────────
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Some headers contain tables (e.g., name + contact in columns)
        for table in section.header.tables:
            for row in table.rows:
                seen = set()
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in seen:
                        seen.add(cell_text)
                        parts.append(cell_text)

    # ── 4. Extract from paragraphs (main body text) ───────────────────────────
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # ── 5. Extract from tables (common in structured resume layouts) ──────────
    # Many resumes use tables for: skills grid, experience timeline, education
    for table in doc.tables:
        for row in table.rows:
            # Deduplicate merged cells (same text repeated across merged columns)
            seen = set()
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in seen:
                    seen.add(cell_text)
                    parts.append(cell_text)

    # ── 6. Extract from document footers ──────────────────────────────────────
    for section in doc.sections:
        for para in section.footer.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

    result = "\n".join(parts)
    logger.debug(f"  DOCX extraction complete: {len(result)} chars from {path.name}")
    return result
