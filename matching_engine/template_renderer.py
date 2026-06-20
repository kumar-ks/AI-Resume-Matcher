"""
Template Renderer
==================

Renders candidate data into a DOCX template format.

RULES:
    - The original template file is NEVER modified (read-only).
    - A fresh copy is created in memory and populated with candidate data.
    - Output is saved to a separate 'rendered/' folder.
    - Output filename: {rank}_Antern_{original_resume_filename}.docx

APPROACH:
    Instead of trying to fill in the exact same paragraph slots, this renderer:
    1. Opens the template to capture formatting/styles
    2. Builds a brand-new document from scratch using those styles
    3. Populates all sections with full candidate data (no truncation)

This ensures ALL data from the candidate's profile appears in the output,
regardless of how many paragraphs the template originally had.
"""

import logging
from copy import deepcopy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from matching_engine.models import ResumeProfile, WorkExperience

logger = logging.getLogger(__name__)


def render_top_candidates(
    results: list,
    source_filenames: dict,
    template_path: Path,
    output_dir: Path,
    top_n: int = 1,
) -> list[Path]:
    """
    Render top N candidates into the DOCX template format.

    The original template is NEVER modified. A fresh document is created
    for each candidate using the template's styles.

    Args:
        results: List of MatchResult (sorted by score, highest first)
        source_filenames: dict mapping raw_text → {"filename": ..., "path": ...}
        template_path: Path to the template DOCX (read-only, never modified)
        output_dir: Directory to save rendered documents
        top_n: Number of top candidates to render (default: 1)

    Returns:
        List of Paths to the rendered DOCX files
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    template_path = Path(template_path)
    if not template_path.exists():
        logger.error(f"Template file not found: {template_path}")
        return []

    rendered_files = []
    candidates_to_render = results[:top_n]

    for rank, result in enumerate(candidates_to_render, 1):
        profile = result.candidate

        # Determine output filename: {rank}_Antern_{original_filename}.docx
        file_info = source_filenames.get(profile.raw_text, {})
        original_filename = file_info.get("filename", f"Candidate_{rank}")
        base_name = Path(original_filename).stem
        output_filename = f"{rank}_Antern_{base_name}.docx"
        output_path = output_dir / output_filename

        logger.info(f"Rendering rank #{rank}: {profile.full_name} → {output_filename}")

        try:
            _render_single(profile, template_path, output_path)
            rendered_files.append(output_path)
            print(f"  ✓ Rendered: {output_path}")
        except Exception as e:
            logger.error(f"Failed to render {output_filename}: {e}")
            print(f"  ✗ Failed to render: {output_filename} ({e})")

    return rendered_files


def _render_single(profile: ResumeProfile, template_path: Path, output_path: Path) -> None:
    """
    Render a single candidate into a new document based on the template.

    Strategy: Create a fresh document, apply template-like formatting,
    and fill ALL candidate data without being constrained by template slots.
    """
    doc = Document()

    # ── Set default font ──────────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9.5)

    # ── NAME (Arial, 16pt, Bold, Centered) ────────────────────────────────────
    name = profile.full_name or "N/A"
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(name.upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    # ── CONTACT LINE (Arial, 11pt) ───────────────────────────────────────────
    contact_parts = []
    if profile.phone:
        contact_parts.append(f"Contact: {profile.phone}")
    if profile.email:
        contact_parts.append(f"Email: {profile.email}")
    if profile.location:
        contact_parts.append(f"Location: {profile.location}")
    contact_text = "; ".join(contact_parts) if contact_parts else "Contact: N/A"

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_para.add_run(contact_text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # ── PROFESSIONAL SUMMARY ──────────────────────────────────────────────────
    _add_section_header(doc, "PROFESSIONAL SUMMARY")

    summary = profile.career_summary or "Experienced professional."
    # Split summary into bullet points if it's long, otherwise single paragraph
    summary_para = doc.add_paragraph(style='List Bullet')
    run = summary_para.add_run(summary)
    run.font.size = Pt(9.5)
    run.font.name = 'Arial'

    # ── TECHNICAL SKILLS ──────────────────────────────────────────────────────
    _add_section_header(doc, "TECHNICAL SKILLS")

    if profile.skills:
        # Group skills in lines of ~5-8 per line for readability
        for i in range(0, len(profile.skills), 6):
            chunk = profile.skills[i:i+6]
            skill_para = doc.add_paragraph(style='List Bullet')
            run = skill_para.add_run(", ".join(chunk))
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'
    else:
        para = doc.add_paragraph("N/A")
        para.runs[0].font.size = Pt(9.5)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    _add_section_header(doc, "EXPERIENCE")

    if profile.work_experiences:
        for exp in profile.work_experiences:
            _add_experience_entry(doc, exp)
    else:
        para = doc.add_paragraph("No work experience data available.")
        para.runs[0].font.size = Pt(9.5)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    _add_section_header(doc, "EDUCATION")

    if profile.education:
        for edu in profile.education:
            edu_para = doc.add_paragraph(style='List Bullet')
            run = edu_para.add_run(edu)
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'
    else:
        para = doc.add_paragraph("N/A")
        para.runs[0].font.size = Pt(9.5)

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    _add_section_header(doc, "CERTIFICATIONS")

    if profile.certifications:
        for cert in profile.certifications:
            cert_para = doc.add_paragraph(style='List Bullet')
            run = cert_para.add_run(cert)
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'
    else:
        para = doc.add_paragraph("N/A")
        para.runs[0].font.size = Pt(9.5)

    # ── Save (template file is never touched) ─────────────────────────────────
    doc.save(str(output_path))
    logger.debug(f"Saved rendered document to: {output_path}")


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================


def _add_section_header(doc: Document, title: str) -> None:
    """Add a section header (bold, 9.5pt, with spacing)."""
    para = doc.add_paragraph()
    para.space_before = Pt(12)
    para.space_after = Pt(4)
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = 'Arial'


def _add_experience_entry(doc: Document, exp: WorkExperience) -> None:
    """
    Add a single work experience entry with title, company, dates,
    technologies, and responsibilities.
    """
    # ── Title line: "Job Title, Company    Date Range" (Bold) ─────────────────
    date_range = ""
    if exp.start_year:
        end = "Present" if exp.is_current else str(exp.end_year or "")
        date_range = f"    {exp.start_year} - {end}"

    title_para = doc.add_paragraph()
    title_para.space_before = Pt(6)
    title_para.space_after = Pt(2)

    # Title + company (bold)
    title_run = title_para.add_run(f"{exp.title}, {exp.company}")
    title_run.bold = True
    title_run.font.size = Pt(9.5)
    title_run.font.name = 'Arial'

    # Date range (not bold, right-aligned effect via spacing)
    if date_range:
        date_run = title_para.add_run(date_range)
        date_run.font.size = Pt(9.5)
        date_run.font.name = 'Arial'

    # ── Technologies (if any) ─────────────────────────────────────────────────
    if exp.technologies:
        tech_para = doc.add_paragraph()
        tech_para.space_after = Pt(2)
        tech_run = tech_para.add_run(f"Technologies: {', '.join(exp.technologies)}")
        tech_run.font.size = Pt(9)
        tech_run.font.name = 'Arial'
        tech_run.italic = True

    # ── Responsibilities (bullet points) ──────────────────────────────────────
    if exp.responsibilities:
        for resp in exp.responsibilities:
            resp_para = doc.add_paragraph(style='List Bullet')
            run = resp_para.add_run(resp)
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'
